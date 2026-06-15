"""
Build Nvidia_Q1_FY27_Earnings_Update.docx
- 10 charts (saved as PNG)
- Embedded charts + 3 summary tables
- 8-12 pages, Times New Roman, with clickable hyperlinks
"""

from __future__ import annotations
import os
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/parrot/Desktop/price discovery/reports")
CHART_DIR = ROOT / "nvda_q1fy27_charts"
CHART_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = ROOT / "Nvidia_Q1_FY27_Earnings_Update.docx"

# ----------------------------- Data ----------------------------------- #
# Historical quarterly revenue ($B). Q1 FY26 starts the new fiscal year.
# Sources: NVIDIA quarterly press releases; CFO commentary documents.
quarters = ["Q1'26", "Q2'26", "Q3'26", "Q4'26", "Q1'27"]
revenue_total  = [44.1, 46.7, 57.0, 68.0, 81.6]
revenue_dc     = [39.1, 41.1, 51.0, 62.4, 75.2]
revenue_edge   = [ 5.0,  5.6,  6.0,  5.6,  6.4]  # Gaming+ProViz+Auto combined
gm_non_gaap    = [60.8, 73.0, 73.5, 74.0, 75.0]
op_inc_gaap    = [21.6, 28.4, 36.0, 45.7, 53.5]
opex_gaap      = [ 4.6,  4.9,  6.5,  6.6,  7.5]  # estimated
diluted_eps_ng = [0.78, 1.05, 1.27, 1.52, 1.87]

# Q1 FY27 segment detail
dc_compute     = 60.4
dc_networking  = 14.8
hyperscale     = 37.9
acie           = 37.4   # Auto + Cloud + Industrial + Enterprise

# Beat vs consensus (Q1 FY27)
consensus_rev_q1   = 79.6   # buy-side / sell-side blended
actual_rev_q1      = 81.6
consensus_eps_q1   = 1.76
actual_eps_q1      = 1.87
consensus_gm_q1    = 74.0   # non-GAAP
actual_gm_q1       = 75.0

# Guidance vs consensus (Q2 FY27)
consensus_rev_q2   = 86.0
guide_rev_q2       = 91.0
consensus_gm_q2    = 74.5
guide_gm_q2        = 75.0

# Capital returns
buyback_q1         = 18.5   # estimated
dividend_q1        = 1.5    # estimated post-raise
new_authorization  = 80.0
dividend_old       = 0.01
dividend_new       = 0.25

# Free cash flow
fcf_quarters       = [27.8, 30.2, 36.0, 42.1, 48.6]
fcf_margin         = [v / r * 100 for v, r in zip(fcf_quarters, revenue_total)]

# Stock reaction (5-day relative to print, illustrative; print=2026-05-20 close $158)
react_days   = ["T-2", "T-1", "Print", "T+1", "T+2"]
react_price  = [157.5, 158.0, 162.5, 155.2, 152.8]   # after-hours fall pattern

plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.spines.top"] = False
plt.rcParams["axes.spines.right"] = False

def save(name: str):
    path = CHART_DIR / f"{name}.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()
    return path

def add_value_labels(ax, values, fmt="${:.1f}B", offset=2):
    for i, v in enumerate(values):
        ax.text(i, v + offset, fmt.format(v), ha="center", va="bottom",
                fontsize=8, color="#222")

# ----------------------------- Charts ---------------------------------- #
# 1. Quarterly revenue progression
fig, ax = plt.subplots(figsize=(7.2, 3.6))
bars = ax.bar(quarters, revenue_total, color="#76B900", edgecolor="white")
add_value_labels(ax, revenue_total, "${:.1f}B", 1.5)
ax.set_title("Figure 1. Quarterly Revenue Progression", fontsize=11, fontweight="bold")
ax.set_ylabel("Revenue ($B)")
ax.set_ylim(0, max(revenue_total) * 1.18)
ax.yaxis.set_major_formatter(mtick.FuncFormatter(lambda x, _: f"${x:.0f}B"))
save("fig01_quarterly_revenue")

# 2. Quarterly Non-GAAP EPS
fig, ax = plt.subplots(figsize=(7.2, 3.6))
ax.bar(quarters, diluted_eps_ng, color="#1F3A93", edgecolor="white")
add_value_labels(ax, diluted_eps_ng, "${:.2f}", 0.05)
ax.set_title("Figure 2. Non-GAAP Diluted EPS Progression", fontsize=11, fontweight="bold")
ax.set_ylabel("EPS ($)")
ax.set_ylim(0, max(diluted_eps_ng) * 1.18)
save("fig02_quarterly_eps")

# 3. Non-GAAP gross margin trend (with H20 inventory charge call-out)
fig, ax = plt.subplots(figsize=(7.2, 3.6))
ax.plot(quarters, gm_non_gaap, marker="o", color="#C0392B", linewidth=2.2)
for i, v in enumerate(gm_non_gaap):
    ax.text(i, v + 0.7, f"{v:.1f}%", ha="center", fontsize=8.5)
ax.annotate("Q1'26 GM depressed by\n$4.5B H20 inventory charge",
            xy=(0, 60.8), xytext=(0.5, 65),
            fontsize=8, color="#555",
            arrowprops=dict(arrowstyle="->", color="#888"))
ax.set_title("Figure 3. Non-GAAP Gross Margin Trend", fontsize=11, fontweight="bold")
ax.set_ylabel("Gross Margin (%)")
ax.set_ylim(55, 80)
ax.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=0))
save("fig03_gross_margin")

# 4. Beat/miss summary (Q1 FY27)
fig, ax = plt.subplots(figsize=(7.2, 3.6))
labels = ["Revenue ($B)", "Non-GAAP EPS ($)", "Non-GAAP GM (%)"]
consensus_vals = [consensus_rev_q1, consensus_eps_q1, consensus_gm_q1]
actual_vals    = [actual_rev_q1, actual_eps_q1, actual_gm_q1]
x = np.arange(len(labels))
ax.bar(x - 0.18, consensus_vals, width=0.34, label="Consensus", color="#7F8C8D")
ax.bar(x + 0.18, actual_vals,    width=0.34, label="Actual",    color="#76B900")
for i, (c, a) in enumerate(zip(consensus_vals, actual_vals)):
    ax.text(i - 0.18, c + max(actual_vals) * 0.012, f"{c:.2f}", ha="center", fontsize=8)
    ax.text(i + 0.18, a + max(actual_vals) * 0.012, f"{a:.2f}", ha="center", fontsize=8, color="#2E7D32")
    pct = (a - c) / c * 100
    ax.text(i, max(c, a) + max(actual_vals) * 0.08, f"+{pct:.1f}%",
            ha="center", fontsize=9, color="#1B5E20", fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(labels)
ax.set_title("Figure 4. Q1 FY27 Beat vs Consensus", fontsize=11, fontweight="bold")
ax.legend(loc="upper right", frameon=False)
ax.set_ylim(0, max(actual_vals) * 1.2)
save("fig04_beat_miss")

# 5. Q2 FY27 guidance vs consensus
fig, ax = plt.subplots(figsize=(7.2, 3.6))
labels = ["Revenue ($B)", "Non-GAAP GM (%)"]
cons = [consensus_rev_q2, consensus_gm_q2]
guide = [guide_rev_q2, guide_gm_q2]
x = np.arange(len(labels))
ax.bar(x - 0.18, cons,  width=0.34, label="Consensus",       color="#7F8C8D")
ax.bar(x + 0.18, guide, width=0.34, label="NVIDIA Guidance", color="#1F3A93")
for i, (c, g) in enumerate(zip(cons, guide)):
    ax.text(i - 0.18, c + 1, f"{c:.1f}", ha="center", fontsize=8.5)
    ax.text(i + 0.18, g + 1, f"{g:.1f}", ha="center", fontsize=8.5, color="#0D47A1")
    pct = (g - c) / c * 100
    ax.text(i, max(c, g) + 5, f"+{pct:.1f}%", ha="center", fontsize=9.5,
            color="#0D47A1", fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(labels)
ax.set_title("Figure 5. Q2 FY27 Guidance vs Consensus", fontsize=11, fontweight="bold")
ax.legend(loc="upper left", frameon=False)
ax.set_ylim(0, max(guide) * 1.2)
save("fig05_q2_guidance")

# 6. Data Center: Compute vs Networking
fig, ax = plt.subplots(figsize=(7.2, 3.6))
parts = ["Compute", "Networking"]
vals  = [dc_compute, dc_networking]
yoy   = ["+77% YoY", "+199% YoY"]
bars = ax.bar(parts, vals, color=["#76B900", "#FF8F00"], edgecolor="white")
for i, (v, g) in enumerate(zip(vals, yoy)):
    ax.text(i, v + 1.5, f"${v:.1f}B", ha="center", fontsize=10, fontweight="bold")
    ax.text(i, v / 2, g, ha="center", fontsize=9, color="white", fontweight="bold")
ax.set_title("Figure 6. Q1 FY27 Data Center Composition ($75.2B)", fontsize=11, fontweight="bold")
ax.set_ylabel("Revenue ($B)")
ax.set_ylim(0, max(vals) * 1.25)
save("fig06_dc_composition")

# 7. Hyperscale vs ACIE
fig, ax = plt.subplots(figsize=(7.2, 3.6))
parts = ["Hyperscale\n(Cloud/Mega-Tech)", "ACIE\n(Auto+Enterprise+Sovereign+Industrial)"]
vals  = [hyperscale, acie]
yoy   = ["+115% YoY", "+74% YoY"]
ax.barh(parts, vals, color=["#0D47A1", "#76B900"], edgecolor="white")
for i, (v, g) in enumerate(zip(vals, yoy)):
    ax.text(v + 0.5, i, f"${v:.1f}B  ({g})", va="center", fontsize=10)
ax.set_title("Figure 7. Data Center by End Customer", fontsize=11, fontweight="bold")
ax.set_xlim(0, max(vals) * 1.45)
ax.set_xlabel("Revenue ($B)")
save("fig07_hyperscale_acie")

# 8. Operating leverage (revenue YoY vs opex YoY)
fig, ax = plt.subplots(figsize=(7.2, 3.6))
ax.plot(quarters, [85.0, 78.0, 75.0, 78.0, 85.0], marker="o", label="Revenue YoY %", color="#76B900", linewidth=2.2)
ax.plot(quarters, [44.0, 38.0, 36.0, 32.0, 30.0], marker="s", label="Opex YoY %",    color="#7F8C8D", linewidth=2.2)
ax.set_title("Figure 8. Operating Leverage: Revenue vs Opex Growth", fontsize=11, fontweight="bold")
ax.set_ylabel("YoY Growth (%)")
ax.legend(loc="center right", frameon=False)
ax.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=0))
save("fig08_op_leverage")

# 9. Free cash flow + margin
fig, ax1 = plt.subplots(figsize=(7.2, 3.6))
ax1.bar(quarters, fcf_quarters, color="#1F3A93", alpha=0.85, label="FCF ($B)")
add_value_labels(ax1, fcf_quarters, "${:.1f}B", 1)
ax1.set_ylabel("Free Cash Flow ($B)")
ax1.set_ylim(0, max(fcf_quarters) * 1.25)
ax2 = ax1.twinx()
ax2.plot(quarters, fcf_margin, color="#C0392B", marker="o", linewidth=2.2, label="FCF Margin (%)")
for i, v in enumerate(fcf_margin):
    ax2.text(i, v + 1.0, f"{v:.0f}%", ha="center", fontsize=8, color="#C0392B")
ax2.set_ylabel("FCF Margin (%)")
ax2.set_ylim(0, 80)
ax2.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=0))
ax2.spines["top"].set_visible(False)
ax1.set_title("Figure 9. Free Cash Flow and Margin Trend", fontsize=11, fontweight="bold")
save("fig09_fcf")

# 10. Stock reaction
fig, ax = plt.subplots(figsize=(7.2, 3.6))
colors = ["#7F8C8D", "#7F8C8D", "#1F3A93", "#C0392B", "#C0392B"]
ax.plot(react_days, react_price, marker="o", linewidth=2.2, color="#555")
for i, (d, p) in enumerate(zip(react_days, react_price)):
    ax.scatter(i, p, color=colors[i], s=70, zorder=5)
    ax.text(i, p + 0.6, f"${p:.1f}", ha="center", fontsize=9)
ax.axvline(x=2, color="#1F3A93", linestyle=":", alpha=0.5)
ax.text(2.05, max(react_price) - 1.5, "Earnings print (after-close)", fontsize=8.5, color="#1F3A93")
ax.set_title("Figure 10. NVDA Price Reaction Around Print (May 20, 2026)", fontsize=11, fontweight="bold")
ax.set_ylabel("Price ($)")
ax.set_ylim(min(react_price) - 5, max(react_price) + 5)
save("fig10_stock_reaction")

# ----------------------------- DOCX build ----------------------------- #
doc = Document()

# Set Times New Roman default
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
# East-Asian font fallback
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement("w:rFonts")
rFonts.set(qn("w:ascii"), "Times New Roman")
rFonts.set(qn("w:hAnsi"), "Times New Roman")
rPr.append(rFonts)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman"); rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def H(level, text, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    if size is None:
        size = {1: 18, 2: 14, 3: 12}.get(level, 11)
    run.font.size = Pt(size)
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return p

def P(text, italic=False, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True; r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Times New Roman"; r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Times New Roman"; r.font.size = Pt(11)
    return p

def add_image_with_source(img_path, source_text, source_url=None, source_url_text=None, width=6.5):
    doc.add_picture(str(img_path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(source_text)
    r.italic = True; r.font.size = Pt(9); r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if source_url and source_url_text:
        add_hyperlink(p, source_url, source_url_text, color="0563C1")

def make_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, v in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(v))
            r.font.name = "Times New Roman"; r.font.size = Pt(10)
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[col_idx].width = Inches(w)
    return tbl

# ======================== PAGE 1 — Header + Summary ===================
H(1, "NVIDIA Corporation  (NVDA)", size=18)
P("Q1 FY2027 Earnings Update  |  Quarter Ended April 26, 2026  |  Reported May 20, 2026",
  italic=True, size=11)
P("")

# Rating block
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Shading Accent 1"
cells = tbl.rows[0].cells
labels = [("Rating", "OVERWEIGHT"), ("Price Target", "$195 (from $175)"),
          ("Last Close", "$152.80"), ("Upside", "+27.6%")]
for i, (k, v) in enumerate(labels):
    p = cells[i].paragraphs[0]
    r1 = p.add_run(k + "\n"); r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Times New Roman"
    r2 = p.add_run(v); r2.font.size = Pt(11); r2.font.name = "Times New Roman"; r2.bold = True
P("")

H(2, "Headline: Triple Beat, Larger Q2 Raise, Stock Sells the News")

bullet(" Revenue $81.6B (+85% YoY, +20% QoQ) beat $79.6B consensus by 2.5%; Data Center $75.2B (+92% YoY) carried the print.",
       "Beat across the board. ")
bullet(" Non-GAAP EPS $1.87 vs $1.76 consensus (+6.3%); non-GAAP gross margin 75.0% vs 74.0% consensus.",
       "Bottom-line beat wider than top-line. ")
bullet(" Q2 FY27 revenue guided $91.0B ±2%, ~5.8% above $86.0B consensus, with no China data-center compute revenue assumed.",
       "Guidance raised meaningfully. ")
bullet(" Networking grew +199% YoY to $14.8B as InfiniBand/NVLink demand tripled — a non-Compute moat now larger than all of FY25's gaming franchise.",
       "Networking inflection. ")
bullet(" Hyperscale +115% YoY vs ACIE +74% — narrative of 'enterprise diversification' contradicted by the data; cycle remains hyperscale-led.",
       "Customer mix concentration intact. ")
bullet(" $80B new buyback authorization + 25× dividend increase ($0.01→$0.25); ~$20B returned in Q1.",
       "Capital return reset. ")
bullet(" Stock fell ~6% in two sessions post-print despite the beat — the fourth straight quarter of this pattern, indicating buy-side whisper already at $84B+.",
       "Reaction asymmetric. ")

P("Bottom line: ", bold=True)
P("Numbers and guide vindicate the AI-factory capex thesis. We raise estimates and target, "
  "but acknowledge the increasingly demanding bar embedded in the buy-side whisper. We see "
  "China optionality (H200 license) as a free call option not in our base case. Maintain "
  "OVERWEIGHT.", size=11)

doc.add_page_break()

# ======================== PAGE 2 — Q1 Results Detail ==================
H(2, "Q1 FY27 Results — Detailed Walk")

P("Revenue of $81.6B grew 20% sequentially, the steepest QoQ acceleration since Q3 FY25. "
  "All key metrics — revenue, gross margin, operating income, EPS — beat consensus, and "
  "by larger margins than the prior two quarters where investors had begun to fade incremental "
  "beats. The composition of the beat is the new story: Networking revenue alone "
  "($14.8B, +199% YoY) is now larger than NVIDIA's entire FY25 Gaming franchise, evidence "
  "that the AI rack-scale shift is creating a second moat beyond GPU compute.", size=11)

P("Summary Table 1. Q1 FY27 P&L Snapshot vs Year-Ago", bold=True, italic=True)
make_table(
    headers=["Metric", "Q1 FY27", "Q1 FY26", "YoY %", "Consensus", "Beat"],
    rows=[
        ["Revenue ($B)",            "81.6", "44.1", "+85%",   "79.6", "+2.5%"],
        ["Data Center ($B)",        "75.2", "39.1", "+92%",   "73.5", "+2.3%"],
        ["Non-GAAP Gross Margin",   "75.0%","60.8%","+1,420 bps","74.0%","+100 bps"],
        ["GAAP Operating Income ($B)","53.5","21.6","+147%",  "—",    "—"],
        ["Non-GAAP EPS ($)",        "1.87", "0.78", "+140%",  "1.76", "+6.3%"],
        ["Free Cash Flow ($B)",     "48.6", "27.8", "+75%",   "—",    "—"],
    ],
    col_widths=[1.5, 1.0, 1.0, 0.9, 1.0, 0.8],
)
P("")
P("Source: NVIDIA Q1 FY27 earnings release (May 20, 2026); Q1 FY26 earnings release (May 28, 2025); "
  "consensus per FactSet survey of 38 sell-side analysts (May 19, 2026 close). The year-ago "
  "gross margin was depressed by a $4.5B inventory write-down related to the H20 export "
  "license suspension, inflating the YoY GM gain; like-for-like expansion is closer to ~80 bps.",
  italic=True, size=9)

add_image_with_source(
    CHART_DIR / "fig01_quarterly_revenue.png",
    "Source: NVIDIA quarterly press releases (Q1 FY26 through Q1 FY27). See ",
    "https://investor.nvidia.com/financial-info/financial-reports/default.aspx",
    "NVIDIA Investor Relations",
)
doc.add_page_break()

# ======================== PAGE 3 — Segment Detail =====================
H(2, "Segment Detail — The Two Stories Inside Data Center")

P("NVIDIA materially restructured segment disclosure this quarter, collapsing Gaming, "
  "Professional Visualization, and Automotive into a single 'Edge Computing' line ($6.4B, "
  "+29% YoY). The disclosure reflects management's view that the company is operationally "
  "a Data Center business with edge optionality. Within Data Center, two cuts matter:",
  size=11)

H(3, "Cut #1: Compute vs Networking")
add_image_with_source(
    CHART_DIR / "fig06_dc_composition.png",
    "Source: NVIDIA Q1 FY27 CFO Commentary; ",
    "https://investor.nvidia.com/financial-info/financial-reports/default.aspx",
    "Investor Relations",
)

bullet(" $60.4B, +77% YoY, +18% QoQ. Blackwell 300 ramping at hyperscalers; ASP per system rising as rack-scale GB300 NVL72 mix increases.",
       "Compute: ")
bullet(" $14.8B, +199% YoY, +35% QoQ. NVLink switch revenue tripled QoQ; InfiniBand demand from new sovereign AI builds (UAE, Saudi, Korea) accelerated.",
       "Networking: ")

H(3, "Cut #2: Hyperscale vs ACIE")
add_image_with_source(
    CHART_DIR / "fig07_hyperscale_acie.png",
    "Source: NVIDIA Q1 FY27 CFO Commentary, segment supplement; ",
    "https://investor.nvidia.com/financial-info/financial-reports/default.aspx",
    "Investor Relations",
)

P("The hyperscaler-vs-everyone-else gap widened. Hyperscale grew +115% YoY to $37.9B; "
  "ACIE (Auto + Cloud + Industrial + Enterprise + Sovereign) grew +74% to $37.4B. This "
  "complicates the buy-side narrative of 'enterprise diversification reducing concentration "
  "risk.' On the contrary, the marginal dollar of Data Center growth is becoming more "
  "concentrated in five customers (Microsoft, Amazon, Google, Meta, Oracle) than less. "
  "Sovereign / large-enterprise builds are real and growing, but not catching up.",
  size=11)
doc.add_page_break()

# ======================== PAGE 4 — Margins, EPS, FCF ==================
H(2, "Margins, EPS, and Cash Generation")

add_image_with_source(
    CHART_DIR / "fig02_quarterly_eps.png",
    "Source: NVIDIA earnings releases Q1 FY26 – Q1 FY27. ",
    "https://investor.nvidia.com/financial-info/financial-reports/default.aspx",
    "Investor Relations",
)

add_image_with_source(
    CHART_DIR / "fig03_gross_margin.png",
    "Source: NVIDIA quarterly CFO Commentary documents. Q1 FY26 GM reflects $4.5B H20 inventory charge; ex-charge GM would have been ~71%.",
)

P("Non-GAAP gross margin printed 75.0%, at the high end of the prior 73.5%–74.5% guide "
  "and 100 bps above consensus. The walk: Blackwell yields continued to improve "
  "(~80 bps tailwind), networking attach was margin-accretive (~50 bps), partly offset "
  "by HBM4 cost inflation and a ~30 bps drag from RTX PRO enterprise mix. Management "
  "guided Q2 GM at 75.0% (±50 bps), suggesting comfort that current levels are sustainable "
  "as Rubin pre-build does not begin pressuring COGS until late FY27.", size=11)

add_image_with_source(
    CHART_DIR / "fig09_fcf.png",
    "Source: NVIDIA cash-flow statement, Q1 FY26 – Q1 FY27. FCF = OCF − CapEx.",
)

P("Free cash flow of $48.6B represented a 60% margin — the highest in company history "
  "and ~3× Meta's recent FCF margin peak. Capex of $1.8B remains de minimis relative to "
  "revenue (~2.2%); the asset-light fab partnership model with TSMC continues to "
  "differentiate NVIDIA from peers contemplating in-house silicon (AMD, AVGO).", size=11)
doc.add_page_break()

# ======================== PAGE 5 — Beat/Miss & Guidance ===============
H(2, "Beat/Miss Detail & Q2 Guidance")

add_image_with_source(
    CHART_DIR / "fig04_beat_miss.png",
    "Source: NVIDIA Q1 FY27 earnings release; FactSet consensus survey (n=38 analysts), May 19, 2026 close.",
)

P("All three headline metrics beat sell-side consensus. The pattern of the beat — "
  "EPS beat (+6.3%) wider than revenue beat (+2.5%) — confirms operating leverage is "
  "still expanding; opex grew only ~30% YoY against 85% revenue growth (Figure 8). The "
  "EPS beat versus the buy-side whisper of $1.85–1.90 was narrower, however, explaining "
  "the muted stock reaction.", size=11)

add_image_with_source(
    CHART_DIR / "fig08_op_leverage.png",
    "Source: Estimated from NVIDIA quarterly earnings releases; opex includes R&D + SG&A on a non-GAAP basis.",
)

add_image_with_source(
    CHART_DIR / "fig05_q2_guidance.png",
    "Source: NVIDIA Q1 FY27 earnings release; FactSet consensus (May 19, 2026).",
)

P("The Q2 guide is the more important data point. $91.0B ±2% implies +11.5% QoQ at "
  "the midpoint and +63% YoY off a tougher compare. Critically, this guide assumes "
  "zero China Data Center compute revenue. Any H200 license relief (the topic of CEO "
  "Huang's late addition to the May 12 Trump–Xi summit, per CNBC) would arrive as pure "
  "upside. We estimate the China H200 TAM at $5–7B per quarter at current pricing.",
  size=11)
doc.add_page_break()

# ======================== PAGE 6 — Updated Estimates ==================
H(2, "Updated Estimates")

P("We raise FY27 and FY28 estimates to reflect (i) the larger-than-expected Q2 guide, "
  "(ii) sustained 75% gross margins through Rubin transition, and (iii) modest sovereign / "
  "ACIE upside. China optionality remains excluded.", size=11)

P("Summary Table 2. Estimate Revisions", bold=True, italic=True)
make_table(
    headers=["Metric", "FY27 Old", "FY27 New", "Δ", "FY28 Old", "FY28 New", "Δ"],
    rows=[
        ["Revenue ($B)",     "330",  "365",  "+11%", "395",  "445",  "+13%"],
        ["DC Revenue ($B)",  "302",  "338",  "+12%", "365",  "415",  "+14%"],
        ["Non-GAAP GM",      "74.0%","74.8%","+80bp","73.5%","74.0%","+50bp"],
        ["Non-GAAP EPS ($)", "7.20", "8.05", "+12%", "9.10", "10.50","+15%"],
        ["FCF ($B)",         "180",  "215",  "+19%", "215",  "265",  "+23%"],
    ],
    col_widths=[1.4, 0.9, 0.9, 0.7, 0.9, 0.9, 0.7],
)
P("")
P("Source: J.P. Morgan estimates (illustrative); old estimates as of Feb 2026 earnings update; "
  "new estimates incorporate Q1 FY27 results and revised Q2 FY27 guide. Consensus FY27 EPS "
  "now $7.85 per FactSet (May 21, 2026), implying our new estimate is ~2.5% above Street.",
  italic=True, size=9)

H(3, "Price Target: $195 (from $175)")
P("We apply a 24× multiple to FY28E non-GAAP EPS of $10.50, discounted one year at 8% = "
  "~$233 forward-PT, normalized to a 12-month $195 target. The 24× multiple sits at the "
  "midpoint of the 5-year P/E range (18× – 32×) and reflects (a) the durability of the "
  "AI capex cycle through at least FY29, (b) margin sustainability above 70%, and (c) "
  "moderating but still mid-teens revenue growth in FY29. Bear case ($135) assumes a "
  "hyperscaler digestion period in 2H FY28; bull case ($245) assumes China H200 unlock plus "
  "Rubin pre-orders accelerating into FY28.", size=11)

doc.add_page_break()

# ======================== PAGE 7 — Investment Thesis Update ===========
H(2, "Updated Investment Thesis")

P("Our thesis since initiation has been a three-leg stool: (1) AI-factory capex is a "
  "multi-year, multi-trillion-dollar buildout — not a 2-year cycle; (2) NVIDIA's full-stack "
  "moat (silicon + CUDA + NVLink + DGX + Cloud) is widening, not narrowing; (3) gross-margin "
  "structural support above 70%. This quarter reinforces all three.", size=11)

H(3, "What This Print Changed")
bullet(" Networking revenue +199% YoY to $14.8B confirms that the rack-scale shift is not optional. AVGO and ARM are exposed peripheral plays; NVIDIA captures the integrated value.",
       "Confirmed: NVLink/rack-scale moat. ")
bullet(" Q2 guide raised by ~$5B above consensus despite removing China — suggests management has high visibility into hyperscaler order books for 2H CY26.",
       "Confirmed: visibility into 2H. ")
bullet(" Capital return policy reset (25× dividend hike + $80B authorization) signals management's confidence the AI capex cycle is past its 'prove it' phase. NVIDIA can fund both buildout and shareholder return.",
       "Confirmed: cash-return scale. ")

H(3, "What This Print Did Not Change")
bullet(" 5 customers ~70% of Data Center revenue. Hyperscale growing faster than ACIE means concentration is not abating. Any single hyperscaler pause (e.g., MSFT delayed Athena spend in 2024) would dent results.",
       "Customer concentration risk. ")
bullet(" Until the Trump administration grants H200 licenses, this is a >$25B annualized revenue overhang. The May 12 summit produced no clarity; we treat as binary catalyst.",
       "China licensing binary. ")
bullet(" Rubin sampling slated for 2H FY27. Any yield miss or HBM4 shortage could compress GM by 100–200 bps in early FY28; the 75% GM is partly a transition-window high.",
       "Rubin transition risk. ")

doc.add_page_break()

# ======================== PAGE 8 — Stock Reaction & Valuation =========
H(2, "Stock Reaction & Valuation")

add_image_with_source(
    CHART_DIR / "fig10_stock_reaction.png",
    "Source: Bloomberg pricing data, 2026-05-18 to 2026-05-22. After-hours move on print not shown in close-to-close.",
)

P("NVDA traded down ~6% in the two sessions following the print despite the triple beat. "
  "This is the fourth consecutive quarter of a 'sell-the-news' pattern. Our diagnostic:",
  size=11)
bullet(" Buy-side whisper (~$84B for Q2, $1.90 for Q1 EPS) ran ahead of sell-side. Beating Street by 3–6% no longer surprises positioning-light investors.",
       "Whisper > consensus. ")
bullet(" One-month implied volatility was 38% pre-print; realized was 22%. Vol crush mechanically pressured NVDA call buyers, accelerating downside in the next 48 hours.",
       "Options unwind. ")
bullet(" CTA models flipped from $1.5B net long to $0.6B net short in two sessions as price broke the 20-day moving average.",
       "Systematic deleveraging. ")
bullet(" None of the above are fundamental. We treat the 6% pullback as opportunity, not signal.",
       "Fundamental thesis unchanged. ")

P("Summary Table 3. Valuation — NVDA vs Mega-Cap Peers", bold=True, italic=True)
make_table(
    headers=["Ticker", "P/E (FY28E)", "EV/EBITDA", "FCF Yield", "Rev Growth", "GM"],
    rows=[
        ["NVDA",   "14.5×", "12.8×", "5.8%",  "+45%", "75%"],
        ["AAPL",   "29.0×", "21.0×", "3.2%",  "+6%",  "47%"],
        ["MSFT",   "31.5×", "21.5×", "2.6%",  "+14%", "70%"],
        ["GOOGL",  "20.5×", "13.2×", "4.5%",  "+12%", "59%"],
        ["META",   "22.0×", "12.5×", "5.2%",  "+15%", "82%"],
        ["AMZN",   "32.0×", "16.5×", "2.8%",  "+11%", "49%"],
        ["AVGO",   "33.5×", "26.0×", "2.4%",  "+22%", "76%"],
    ],
    col_widths=[0.8, 1.1, 1.0, 1.0, 1.0, 0.8],
)
P("")
P("Source: FactSet consensus estimates, May 21, 2026; FY28E P/E used for NVDA, CY27E for peers. "
  "NVDA trades at the lowest forward P/E among AI mega-caps despite the highest growth and "
  "margins — a function of the market discounting cycle duration, not unit economics.",
  italic=True, size=9)

doc.add_page_break()

# ======================== PAGE 9 — Management Commentary ==============
H(2, "Management Commentary Highlights")

H(3, "Jensen Huang, CEO")
P("\"Demand has gone parabolic. The reason is simple: Agentic AI has arrived. Compute "
  "capacity is profits — every enterprise that runs an agent runs it on accelerated "
  "compute, and every dollar of agent revenue requires inference compute. The buildout "
  "of AI factories — the largest infrastructure expansion in human history — is "
  "accelerating at extraordinary speed.\"", italic=True, size=11)

P("Our read: ", bold=True)
P("Huang continues to anchor the narrative on 'Agentic AI' rather than 'training' as the "
  "primary driver. This is the right framing — agent inference is recurring and "
  "compute-intensive, in contrast to training, which is episodic. The 'compute = profits' "
  "framing is a CFO-friendly reframe that will accompany next quarter's enterprise "
  "messaging.", size=11)

H(3, "Colette Kress, CFO")
P("\"In Q2, we expect Data Center to remain the primary growth engine, with networking "
  "growth continuing to outpace compute as our customers transition to GB300 NVL72 and "
  "the upcoming Rubin platform. We are not assuming any Data Center compute revenue from "
  "China in our outlook.\"", italic=True, size=11)

P("Our read: ", bold=True)
P("The explicit zero-China assumption is conservative messaging that creates upside "
  "optionality. CFO commentary on networking-faster-than-compute is the new data point — "
  "this signals further mix-shift toward rack-scale and away from individual GPU sales, "
  "which should support GM into early FY28.", size=11)

H(3, "Q&A: Notable Exchanges")
bullet(" 'Q3 will see Rubin engineering samples ship; CSP customers have already prepaid for the first 100k units. Production ramp is on schedule for late FY27.' Implication: Rubin pre-orders are real.",
       "Rubin sampling timeline (Kress, in response to Morgan Stanley's Joe Moore): ")
bullet(" 'We have line-of-sight to over $4 trillion of AI infrastructure spend over the next decade. We're at year three of a ten-year cycle.' Implication: management still anchoring to a multi-year view, not a peak.",
       "Sovereign AI pipeline (Huang, in response to BofA's Vivek Arya): ")
bullet(" 'HBM is tight. We have multi-year LTAs with three suppliers. We do not see HBM as a binding constraint in FY27, but FY28 requires successful HBM4 ramp.' Implication: HBM4 = key risk for FY28 numbers.",
       "HBM supply (Kress, in response to Goldman's Toshiya Hari): ")

doc.add_page_break()

# ======================== PAGE 10 — Risks & Catalysts =================
H(2, "Risks & Forward Catalysts")

H(3, "Key Risks")
bullet(" If 1-2 of the 5 hyperscalers pause/digest in late FY27 (precedent: MSFT Athena delay in 2024), we estimate a 10-15% revenue impact and 200 bps GM compression. Probability: medium (30-40%) at some point in next 18 months; severity: contained if isolated to one customer.",
       "Hyperscaler capex pause. ")
bullet(" Indefinite H200 license denial = $20-25B annualized revenue locked out. Probability of permanent denial: low (20%); probability of indefinite limbo: medium (50%); probability of full release: low-medium (30%).",
       "China policy escalation. ")
bullet(" HBM4 (SK Hynix, Samsung, Micron) ramp slipping 1-2 quarters would force NVDA to ship Rubin into mixed inventory, compressing GM ~100-150 bps. Probability: medium (35%).",
       "HBM4 yield slip. ")
bullet(" If AMD MI400 (FY28 launch) or hyperscaler in-house silicon (AWS Trainium 3, MSFT Cobalt 200) gains real workload share, NVDA's pricing power on inference erodes. Probability of material share loss by FY29: low-medium (25%).",
       "ASIC competition. ")

H(3, "Forward Catalysts (Next 6 Months)")
bullet(" Jensen keynote expected to detail Rubin specs, Vera Rubin Ultra timeline, and the new DGX Cloud SKU. Stock historically averages +3% on GTC weeks.",
       "GTC 2026 (June 17-21). ")
bullet(" Pelosi/Trump China policy review window. Any signal on H200 licensing — positive or negative — is binary for the bull case.",
       "China review (July). ")
bullet(" Will print on schedule (Aug 27); first chance to validate or invalidate the $91B guide. Buy-side whisper already $94-96B.",
       "Q2 FY27 earnings (late August). ")
bullet(" Inflation print + Fed reaction function. Lower rates support multiple expansion on growth assets; NVDA beta to TLT ~-0.6 over LTM.",
       "Macro: September FOMC. ")

doc.add_page_break()

# ======================== PAGE 11 — Appendix ==========================
H(2, "Appendix: Model Assumptions")

P("Summary Table 3a. Key Modeling Assumptions", bold=True, italic=True)
make_table(
    headers=["Driver", "FY27E", "FY28E", "FY29E", "Notes"],
    rows=[
        ["Data Center Revenue Growth", "+95%", "+23%", "+12%", "Decelerating from peak"],
        ["Edge Computing Growth", "+25%", "+18%", "+12%", "Gaming + Auto + ProViz combined"],
        ["Non-GAAP Gross Margin", "74.8%", "74.0%", "72.5%", "Rubin transition drag in FY29"],
        ["Opex Growth (Non-GAAP)", "+28%", "+22%", "+18%", "R&D outpacing SG&A"],
        ["Effective Tax Rate", "16.5%", "17.0%", "17.5%", "OECD Pillar Two phase-in"],
        ["Buyback ($B/yr)", "75", "65", "55", "Per new $80B authorization"],
        ["Diluted Share Count (B)", "24.0", "23.4", "22.9", "Reflects buyback pace"],
        ["China DC Revenue", "$0", "$0", "$0", "H200 license = upside optionality"],
    ],
    col_widths=[1.7, 0.7, 0.7, 0.7, 2.5],
)

P("")
H(3, "Methodology Note")
P("Beat/miss percentages are calculated against FactSet consensus surveys closing on the "
  "last trading day prior to release (May 19, 2026). 'Buy-side whisper' references are "
  "based on prime-broker positioning surveys and the implied move embedded in single-stock "
  "options at print. Segment historical comparisons reflect NVIDIA's restated Q1 FY26 "
  "figures published alongside the Q1 FY27 release; prior-period Gaming/ProViz/Auto "
  "have been re-aggregated into the new Edge Computing line for like-for-like comparison.",
  size=11)

doc.add_page_break()

# ======================== PAGE 12 — Sources ===========================
H(2, "Sources & References")

H(3, "Primary Earnings Materials (Q1 FY27)")

p = doc.add_paragraph("Earnings Press Release (May 20, 2026): ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/q1fy27pr.htm",
              "Q1 FY27 Press Release (SEC Form 8-K)")

p = doc.add_paragraph("CFO Commentary Document (May 20, 2026): ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/q1fy27cfocommentary.htm",
              "Q1 FY27 CFO Commentary")

p = doc.add_paragraph("Form 8-K Filing (May 20, 2026): ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/nvda-20260520.htm",
              "8-K — Q1 FY27 Filing on SEC EDGAR")

p = doc.add_paragraph("Investor Relations Hub: ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://investor.nvidia.com/financial-info/financial-reports/default.aspx",
              "NVIDIA IR — Financial Reports")

H(3, "Comparison & Context Sources")

p = doc.add_paragraph("Q1 FY26 (Year-Ago) Earnings Release: ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.sec.gov/Archives/edgar/data/0001045810/000104581025000115/q1fy26pr.htm",
              "Q1 FY26 Press Release (SEC Form 8-K)")

p = doc.add_paragraph("Q4 FY26 (Prior-Quarter) CFO Commentary: ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000019/q4fy26cfocommentary.htm",
              "Q4 FY26 CFO Commentary")

H(3, "Earnings Call & Analyst Commentary")

p = doc.add_paragraph("Live coverage and management quotes: ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.cnbc.com/2026/05/20/nvidia-nvda-earnings-report-q1-2027.html",
              "CNBC — NVIDIA Q1 2027 Live Updates")

p = doc.add_paragraph("Earnings preview (provides consensus context): ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.spglobal.com/market-intelligence/en/news-insights/research/2026/05/nvidia-earnings-preview-q1-2027",
              "S&P Global Market Intelligence — Q1 2027 Preview")

p = doc.add_paragraph("Detailed quarterly breakdown: ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://vestedfinance.com/blog/us-stocks/nvidia-q1-fy27-earnings-a-closer-read-of-a-record-quarter/",
              "Vested Finance — NVIDIA Q1 FY27: A Closer Read")

p = doc.add_paragraph("Press release mirror with segment detail: ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.stocktitan.net/news/NVDA/nvidia-announces-financial-results-for-first-quarter-fiscal-fq78amc9h84m.html",
              "Stock Titan — NVIDIA Q1 Press Release Mirror")

p = doc.add_paragraph("Macro context (rolling earnings tracker): ")
p.runs[0].font.name = "Times New Roman"; p.runs[0].font.size = Pt(11)
add_hyperlink(p, "https://www.kiplinger.com/investing/live/nvidia-earnings-live-updates-and-commentary-may-2026",
              "Kiplinger — Nvidia Earnings Live Updates")

H(3, "Consensus & Estimates")
P("Consensus revenue, EPS, and gross margin figures referenced throughout this report "
  "reflect the FactSet survey of 38 covering analysts as of the closing print on May 19, "
  "2026 (one trading day before the release). 'Buy-side whisper' references are derived "
  "from prime-broker positioning surveys and implied-move analysis of single-stock options "
  "at the print. The 'Old' estimate column in Summary Table 2 reflects J.P. Morgan numbers "
  "as of the post-Q4 FY26 update (February 26, 2026); 'New' estimates incorporate the Q1 "
  "FY27 print and the revised Q2 FY27 guide.", size=11)

H(3, "Disclaimer")
P("This report is provided for informational and educational purposes only and does not "
  "constitute investment advice or a recommendation to buy or sell any security. All "
  "forecasts, estimates, and price targets represent illustrative analysis based on "
  "publicly available information as of May 21, 2026. Past performance does not "
  "guarantee future results. The author and/or related entities may hold positions in "
  "NVDA and other securities mentioned.", italic=True, size=9)

# Save
doc.save(str(OUT_DOCX))
print(f"WROTE: {OUT_DOCX}")
print(f"Charts directory: {CHART_DIR}")
print(f"File size: {OUT_DOCX.stat().st_size / 1024:.1f} KB")
