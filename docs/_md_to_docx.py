"""
Markdown → DOCX converter (custom for price_discovery_logic.md)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SCRIPT_DIR = Path(__file__).parent
MD_FILE = SCRIPT_DIR / "price_discovery_logic.md"
DOCX_FILE = SCRIPT_DIR / "price_discovery_logic.docx"


def set_cell_background(cell, color_hex):
    """Set cell background color via XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_inline_runs(para, text, base_size=10):
    """Parse inline markdown (bold **text**, code `text`) and add runs."""
    # Pattern: bold **...** or inline code `...`
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.font.size = Pt(base_size)
        match_text = m.group(0)
        if match_text.startswith("**"):
            run = para.add_run(match_text[2:-2])
            run.font.size = Pt(base_size)
            run.bold = True
        elif match_text.startswith("`"):
            run = para.add_run(match_text[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = Pt(base_size)


def parse_table(lines, idx):
    """Parse a markdown table starting at idx. Returns (rows, next_idx)."""
    rows = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.strip("|").split("|")]
        # skip separator row (---)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            idx += 1
            continue
        rows.append(cells)
        idx += 1
    return rows, idx


def add_table(doc, rows, header_row=True):
    """Add a styled table to the document."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad rows to same length
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline_runs(para, cell_text, base_size=9)
            if ri == 0 and header_row:
                set_cell_background(cell, "1F4E79")
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block toggle
        if stripped.startswith("```"):
            if in_code_block:
                # Close code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                for cl in code_lines:
                    run = p.add_run(cl + "\n")
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Background shading via paragraph
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F4F4F4")
                pPr.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headers
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading = doc.add_heading(stripped[2:], level=0)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue
        if stripped.startswith("## "):
            heading = doc.add_heading(stripped[3:], level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue
        if stripped.startswith("### "):
            heading = doc.add_heading(stripped[4:], level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            i += 1
            continue
        if stripped.startswith("#### "):
            heading = doc.add_heading(stripped[5:], level=3)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), "B0B0B0")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Block quote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows, header_row=True)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, content)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(MD_FILE, DOCX_FILE)
