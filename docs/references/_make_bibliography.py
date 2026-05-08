"""
Generate bibliography DOCX (then converted to PDF if pandoc available).
Maps each Price Discovery system component to its academic foundation.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SCRIPT_DIR = Path(__file__).parent
OUTPUT_DOCX = SCRIPT_DIR / "00_Bibliography_PriceDiscovery_References.docx"


# ─── Reference data ─────────────────────────────────────────────────────────

REFERENCES = [
    # ─ Momentum & Cross-Section ─
    {
        "id": "01",
        "filename": "01_Jegadeesh_Titman_1993_Momentum.pdf",
        "title": "Returns to Buying Winners and Selling Losers: Implications for Stock Market Efficiency",
        "authors": "Jegadeesh, N. & Titman, S.",
        "year": "1993",
        "venue": "Journal of Finance, 48(1), 65-91",
        "applies_to": [
            "Cross-sectional momentum (RSS) — 12-1M return percentile",
            "ret_12_1m field as canonical momentum measure",
            "기본 모멘텀 효과의 학술적 근거",
        ],
        "category": "Momentum (Foundational)",
    },
    {
        "id": "02",
        "filename": "02_Moskowitz_Ooi_Pedersen_2012_TSMomentum.pdf",
        "title": "Time Series Momentum",
        "authors": "Moskowitz, T. J., Ooi, Y. H., & Pedersen, L. H.",
        "year": "2012",
        "venue": "Journal of Financial Economics, 104(2), 228-250",
        "applies_to": [
            "TCS — Trend Continuation Score (단기/장기 추세 지속성)",
            "Multi-asset momentum across ETFs (FICC + Equity)",
            "1-12개월 horizon에서 자기 시계열 추세의 통계적 유의성",
        ],
        "category": "Momentum (Multi-Asset)",
    },
    {
        "id": "03",
        "filename": "03_Asness_Moskowitz_Pedersen_2013_ValMom.pdf",
        "title": "Value and Momentum Everywhere",
        "authors": "Asness, C. S., Moskowitz, T. J., & Pedersen, L. H.",
        "year": "2013",
        "venue": "Journal of Finance, 68(3), 929-985",
        "applies_to": [
            "Multi-asset class momentum (8 markets)",
            "Pre-Momentum/Momentum 구분의 cross-asset validity",
            "Factor framework 일반화",
        ],
        "category": "Momentum (Cross-Asset)",
    },
    {
        "id": "16",
        "filename": "16_Jegadeesh_Titman_2001_MomentumProfitability.pdf",
        "title": "Profitability of Momentum Strategies: An Evaluation of Alternative Explanations",
        "authors": "Jegadeesh, N. & Titman, S.",
        "year": "2001",
        "venue": "Journal of Finance, 56(2), 699-720",
        "applies_to": [
            "1990년대 모멘텀 robustness 검증 — 시스템의 시간대 일반화 근거",
            "behavioral overreaction 모델 검증",
            "long-horizon reversal (4-5년) → CYCLE_PEAK 분류 동기",
        ],
        "category": "Momentum (Robustness)",
    },
    {
        "id": "06",
        "filename": "06_George_Hwang_2004_52WeekHigh.pdf",
        "title": "The 52-Week High and Momentum Investing",
        "authors": "George, T. J. & Hwang, C-Y.",
        "year": "2004",
        "venue": "Journal of Finance, 59(5), 2145-2176",
        "applies_to": [
            "OER (Overextension Risk) — 52주 고점 근접도 사용",
            "pct_from_high 지표의 학술적 정당성",
            "52주 고점 anchoring 효과 → momentum 강화 신호",
        ],
        "category": "Momentum (52-Week High)",
    },
    {
        "id": "05",
        "filename": "05_Daniel_Moskowitz_2016_MomentumCrashes.pdf",
        "title": "Momentum Crashes",
        "authors": "Daniel, K. D. & Moskowitz, T. J.",
        "year": "2016",
        "venue": "Journal of Financial Economics, 122(2), 221-247",
        "applies_to": [
            "OER (Overextension Risk) — 모멘텀 과열 후 급락 위험",
            "OVEREXTENDED override 도입 근거",
            "ret_36_12m + 12M 모멘텀 둔화 → CYCLE_PEAK 분류",
            "vol-managed momentum 아이디어 (Phase 1: vol-adjusted buffer)",
        ],
        "category": "Momentum (Risk Management)",
    },

    # ─ Behavioral / Underreaction ─
    {
        "id": "11",
        "filename": "11_Barberis_Shleifer_Vishny_1998_InvestorSentiment.pdf",
        "title": "A Model of Investor Sentiment",
        "authors": "Barberis, N., Shleifer, A., & Vishny, R.",
        "year": "1998",
        "venue": "Journal of Financial Economics, 49(3), 307-343",
        "applies_to": [
            "Pre-Momentum 4-Agent의 'underreaction' 가정",
            "Microstructure Agent — 변동성 압축, 축적 패턴 (psychological anchoring)",
            "investor sentiment shifts → momentum + reversal 모두 설명",
        ],
        "category": "Behavioral (Underreaction)",
    },
    {
        "id": "04",
        "filename": "04_Hong_Stein_1999_UnifiedTheory.pdf",
        "title": "A Unified Theory of Underreaction, Momentum Trading, and Overreaction in Asset Markets",
        "authors": "Hong, H. & Stein, J. C.",
        "year": "1999",
        "venue": "Journal of Finance, 54(6), 2143-2184",
        "applies_to": [
            "Graph Relational Agent — peer momentum 전파, lead-lag 효과",
            "정보 확산 지연 → 후행 종목 LAGGING_CATCHUP 분류 근거",
            "newswatchers + momentum traders 공존 모델",
        ],
        "category": "Behavioral (Information Diffusion)",
    },
    {
        "id": "12",
        "filename": "12_Daniel_Hirshleifer_Subrahmanyam_1998_InvestorPsychology.pdf",
        "title": "Investor Psychology and Security Market Under- and Overreactions",
        "authors": "Daniel, K. D., Hirshleifer, D., & Subrahmanyam, A.",
        "year": "1998",
        "venue": "Journal of Finance, 53(6), 1839-1885",
        "applies_to": [
            "단기 momentum + 장기 reversal 동시 설명 (overconfidence bias)",
            "OER (Overextension Risk) 설계 — 자기귀인편향 후 mean reversion",
            "Conviction levels (HIGH/MEDIUM/LOW) — 신호 합의도",
        ],
        "category": "Behavioral (Overconfidence)",
    },
    {
        "id": "19",
        "filename": "19_DeBondt_Thaler_1985_StockMarketOverreact.pdf",
        "title": "Does the Stock Market Overreact?",
        "authors": "De Bondt, W. F. M. & Thaler, R.",
        "year": "1985",
        "venue": "Journal of Finance, 40(3), 793-805",
        "applies_to": [
            "Long-term reversal — 36-12M reversal_pctile 사용",
            "CYCLE_PEAK 분류의 행동학적 근거",
            "Loser portfolio 반등 → COUNTER_RALLY/RECOVERY 분류",
        ],
        "category": "Behavioral (Overreaction)",
    },

    # ─ Factor Models ─
    {
        "id": "07",
        "filename": "07_Fama_French_1993_3Factor.pdf",
        "title": "Common Risk Factors in the Returns on Stocks and Bonds",
        "authors": "Fama, E. F. & French, K. R.",
        "year": "1993",
        "venue": "Journal of Financial Economics, 33(1), 3-56",
        "applies_to": [
            "Multi-factor model 기반 (Composite Score 가중 합)",
            "Sector / Size 매핑 → benchmark 대비 excess return",
            "Factor Efficacy 탭의 PCA factor 분석 근거",
        ],
        "category": "Factor Model (Foundational)",
    },
    {
        "id": "08",
        "filename": "08_Fama_MacBeth_1973_RiskReturn.pdf",
        "title": "Risk, Return, and Equilibrium: Empirical Tests",
        "authors": "Fama, E. F. & MacBeth, J. D.",
        "year": "1973",
        "venue": "Journal of Political Economy, 81(3), 607-636",
        "applies_to": [
            "Cross-sectional regression methodology (Factor Efficacy 탭)",
            "Two-pass procedure: time-series β 추정 → cross-sectional λ",
            "Factor premium 통계적 유의성 검정",
        ],
        "category": "Factor Model (Methodology)",
    },
    {
        "id": "09",
        "filename": "09_Carhart_1997_MutualFundPersistence.pdf",
        "title": "On Persistence in Mutual Fund Performance",
        "authors": "Carhart, M. M.",
        "year": "1997",
        "venue": "Journal of Finance, 52(1), 57-82",
        "applies_to": [
            "4-factor model (Mkt + SMB + HML + UMD/Momentum)",
            "Hedge strategy 가중치 — Momentum factor 비중 1.5x (O'Neil)",
            "Performance persistence 분석 (Validation 탭)",
        ],
        "category": "Factor Model (Momentum Factor)",
    },
    {
        "id": "14",
        "filename": "14_Asness_Frazzini_Pedersen_2014_QualityMinusJunk.pdf",
        "title": "Quality Minus Junk",
        "authors": "Asness, C. S., Frazzini, A., & Pedersen, L. H.",
        "year": "2014",
        "venue": "Review of Accounting Studies, 24(1), 34-112",
        "applies_to": [
            "Structural Quality (structural_q) 점수 — 종목별 quality 측정",
            "Pre-Momentum Microstructure: structural_divergence 신호",
            "Quality factor — 안전/수익성/성장/관리 4축",
        ],
        "category": "Factor Model (Quality)",
    },
    {
        "id": "15",
        "filename": "15_Frazzini_Pedersen_2014_BettingAgainstBeta.pdf",
        "title": "Betting Against Beta",
        "authors": "Frazzini, A. & Pedersen, L. H.",
        "year": "2014",
        "venue": "Journal of Financial Economics, 111(1), 1-25",
        "applies_to": [
            "Beta-adjusted strategy — vol-normalized momentum",
            "Pre-Momentum Catalyst: vol_adj_mom 신호",
            "Low-beta anomaly → 시장 레짐 어댑티브 가중치",
        ],
        "category": "Factor Model (BAB)",
    },

    # ─ Technical Analysis ─
    {
        "id": "10",
        "filename": "10_Lo_Mamaysky_Wang_2000_TechnicalAnalysis.pdf",
        "title": "Foundations of Technical Analysis: Computational Algorithms, Statistical Inference, and Empirical Implementation",
        "authors": "Lo, A. W., Mamaysky, H., & Wang, J.",
        "year": "2000",
        "venue": "Journal of Finance, 55(4), 1705-1765",
        "applies_to": [
            "기술적 분석의 학술적 정당성",
            "Pattern recognition (Wyckoff/Darvas/Ichimoku)",
            "8-Hedge Strategy의 학술적 근거",
            "SMA, RSI 등 indicator의 통계적 유의성",
        ],
        "category": "Technical Analysis (Foundational)",
    },
    {
        "id": "13",
        "filename": "13_Brock_Lakonishok_LeBaron_1992_TechnicalRules.pdf",
        "title": "Simple Technical Trading Rules and the Stochastic Properties of Stock Returns",
        "authors": "Brock, W., Lakonishok, J., & LeBaron, B.",
        "year": "1992",
        "venue": "Journal of Finance, 47(5), 1731-1764",
        "applies_to": [
            "Moving Average crossover rules — SMA20/50/200 사용",
            "TCS / TFS 점수의 base 신호 (above_sma50, sma50_slope 등)",
            "Trading range breakout — Darvas Box 전략",
        ],
        "category": "Technical Analysis (MA Rules)",
    },

    # ─ Volatility / Risk ─
    {
        "id": "18",
        "filename": "18_Andersen_Bollerslev_1998_RealizedVolatility.pdf",
        "title": "Answering the Skeptics: Yes, Standard Volatility Models Do Provide Accurate Forecasts",
        "authors": "Andersen, T. G. & Bollerslev, T.",
        "year": "1998",
        "venue": "International Economic Review, 39(4), 885-905",
        "applies_to": [
            "realized_vol 측정 방법론",
            "Phase 1 (Vol-adjusted buffer) — daily σ 계산",
            "vol_3y_ann 칼럼의 학술적 근거",
            "Variance Contraction Pattern (VCP)",
        ],
        "category": "Volatility (Foundational)",
    },

    # ─ Risk-Adjusted Performance ─
    {
        "id": "20",
        "filename": "20_Kidd_SharpeRatio_InfoRatio.pdf",
        "title": "The Sharpe Ratio and the Information Ratio",
        "authors": "Kidd, D. (citing Sharpe 1994)",
        "year": "2011",
        "venue": "CFA Institute Investment Risk and Performance",
        "applies_to": [
            "Sharpe Ratio 산출 — Validation 탭의 risk-adjusted metric",
            "Information Ratio — Grinold's Fundamental Law (IC × √Breadth)",
            "Pre-Momentum Conviction 정량화 근거",
        ],
        "category": "Risk-Adjusted Performance",
    },

    # ─ Option Pricing (참고) ─
    {
        "id": "17",
        "filename": "17_Black_Scholes_1973_OptionPricing.pdf",
        "title": "The Pricing of Options and Corporate Liabilities",
        "authors": "Black, F. & Scholes, M.",
        "year": "1973",
        "venue": "Journal of Political Economy, 81(3), 637-654",
        "applies_to": [
            "변동성 측정의 이론적 근거",
            "Put option hedging — Pipeline 탭의 HEDGE 액션 권고",
            "Implied vol 개념 (향후 확장 가능)",
        ],
        "category": "Option Theory (Reference)",
    },
]

# Papers we couldn't directly download (provide citation only)
UNDOWNLOADED_REFERENCES = [
    {
        "title": "Is Momentum Really Momentum?",
        "authors": "Novy-Marx, R.",
        "year": "2012",
        "venue": "Journal of Financial Economics, 103(3), 429-453",
        "url": "https://rnm.simon.rochester.edu/research/MOM.pdf",
        "applies_to": "Intermediate-horizon momentum (12-7M) → ret_36_12m, ret_12_1m 계산 정당성",
        "note": "직접 다운로드 불가 (저자 사이트 차단). URL 방문 필요.",
    },
    {
        "title": "The Sharpe Ratio (Original)",
        "authors": "Sharpe, W. F.",
        "year": "1994",
        "venue": "Journal of Portfolio Management, 21(1), 49-58",
        "url": "https://web.stanford.edu/~wfsharpe/art/sr/sr.htm",
        "applies_to": "Sharpe Ratio 정의 (web 버전 제공).",
        "note": "JPM 공식 버전은 paywalled — 위 URL의 web summary 참고.",
    },
    {
        "title": "The Fundamental Law of Active Management",
        "authors": "Grinold, R. C.",
        "year": "1989",
        "venue": "Journal of Portfolio Management, 15(3), 30-37",
        "url": "https://jpm.pm-research.com/content/15/3/30",
        "applies_to": "Information Ratio = IC × √Breadth — Pre-Momentum scoring 정당화",
        "note": "Paywalled. AnalystPrep 등 secondary 자료 참고.",
    },
    {
        "title": "New Concepts in Technical Trading Systems (book)",
        "authors": "Wilder, J. W.",
        "year": "1978",
        "venue": "Trend Research",
        "url": "—",
        "applies_to": "RSI (Relative Strength Index) 정의 — Hedge Strategy 입력",
        "note": "도서. 14-day RSI 계산법 원전.",
    },
    {
        "title": "VCP / SEPA",
        "authors": "Minervini, M.",
        "year": "2013",
        "venue": "Trade Like a Stock Market Wizard (book)",
        "url": "—",
        "applies_to": "Hedge Strategy: Minervini SEPA Stage 2 template",
        "note": "도서. 학술 reference 아님.",
    },
    {
        "title": "How to Make Money in Stocks (CANSLIM)",
        "authors": "O'Neil, W. J.",
        "year": "1988",
        "venue": "McGraw-Hill (book)",
        "url": "—",
        "applies_to": "Hedge Strategy: O'Neil CANSLIM 가중치 1.5x",
        "note": "도서. CANSLIM 7-criteria framework.",
    },
]


# ─── DOCX builder helpers ───────────────────────────────────────────────────

def add_inline_runs(para, text, base_size=10):
    import re
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.font.size = Pt(base_size)
        match_text = m.group(0)
        if match_text.startswith("**"):
            run = para.add_run(match_text[2:-2])
            run.font.size = Pt(base_size)
            run.bold = True
        elif match_text.startswith("`"):
            run = para.add_run(match_text[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = Pt(base_size)


def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def make_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Title
    h = doc.add_heading("Price Discovery System — Academic References", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    add_inline_runs(p, "**작성일**: 2026-04-30   |   **수록 논문**: 20개 (다운로드) + 6개 (citation only)")
    doc.add_paragraph()

    # Intro
    doc.add_heading("Overview", level=1)
    intro = doc.add_paragraph()
    add_inline_runs(intro,
        "Price Discovery 시스템에서 사용된 모든 정량 로직은 학술적으로 검증된 연구에 기반합니다. "
        "이 문서는 시스템 컴포넌트별로 그 학술적 근거를 매핑한 reference 모음입니다."
    )
    doc.add_paragraph()

    # System component mapping summary
    doc.add_heading("Component → Reference Mapping", level=1)
    mapping = [
        ("Composite Score (TCS/TFS/RSS)", "Jegadeesh-Titman 1993, Moskowitz 2012, Brock 1992"),
        ("OER (Overextension Risk)", "Daniel-Moskowitz 2016, George-Hwang 2004, DHS 1998"),
        ("Pre-Momentum 4-Agent", "Hong-Stein 1999, BSV 1998, DHS 1998, AFP 2014 (Quality)"),
        ("Hedge Strategies (8개)", "Lo-Mamaysky-Wang 2000, Brock-Lakonishok 1992, O'Neil/Minervini (books)"),
        ("Classification (3×3 + overrides)", "Carhart 1997, Daniel-Moskowitz 2016, Asness 2013"),
        ("Phase 1: Vol-Adjusted Buffer", "Andersen-Bollerslev 1998, Frazzini-Pedersen 2014 (BAB)"),
        ("Validation (Forward Returns / IC / Persistence)", "Grinold 1989, Sharpe 1994, Fama-MacBeth 1973"),
        ("LAGGING_CATCHUP / Underreaction", "Hong-Stein 1999, BSV 1998"),
        ("Long-term Reversal (CYCLE_PEAK)", "De Bondt-Thaler 1985, Jegadeesh-Titman 2001"),
        ("Multi-horizon Returns (1D~5Y)", "Asness 2013, Moskowitz 2012, Carhart 1997"),
    ]
    table = doc.add_table(rows=len(mapping) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    table.rows[0].cells[0].text = "System Component"
    table.rows[0].cells[1].text = "Academic References"
    for c in table.rows[0].cells:
        set_cell_background(c, "1F4E79")
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (comp, refs) in enumerate(mapping):
        table.rows[i + 1].cells[0].text = comp
        table.rows[i + 1].cells[1].text = refs

    doc.add_paragraph()

    # Detailed references by category
    doc.add_heading("References by Category", level=1)

    categories = {}
    for r in REFERENCES:
        categories.setdefault(r["category"], []).append(r)

    for cat in [
        "Momentum (Foundational)",
        "Momentum (Multi-Asset)",
        "Momentum (Cross-Asset)",
        "Momentum (Robustness)",
        "Momentum (52-Week High)",
        "Momentum (Risk Management)",
        "Behavioral (Underreaction)",
        "Behavioral (Information Diffusion)",
        "Behavioral (Overconfidence)",
        "Behavioral (Overreaction)",
        "Factor Model (Foundational)",
        "Factor Model (Methodology)",
        "Factor Model (Momentum Factor)",
        "Factor Model (Quality)",
        "Factor Model (BAB)",
        "Technical Analysis (Foundational)",
        "Technical Analysis (MA Rules)",
        "Volatility (Foundational)",
        "Risk-Adjusted Performance",
        "Option Theory (Reference)",
    ]:
        if cat not in categories:
            continue
        h = doc.add_heading(cat, level=2)
        for r in categories[cat]:
            # Title with id
            p = doc.add_paragraph()
            run = p.add_run(f"[{r['id']}] {r['title']}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            # Authors / year / venue
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.2)
            run = p2.add_run(f"{r['authors']} ({r['year']}). ")
            run.italic = True
            run.font.size = Pt(10)
            run = p2.add_run(r["venue"])
            run.font.size = Pt(10)

            # Filename
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Inches(0.2)
            run = p3.add_run(f"📄 {r['filename']}")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)

            # Applies to
            p4 = doc.add_paragraph()
            p4.paragraph_format.left_indent = Inches(0.2)
            run = p4.add_run("Applies to:")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            for item in r["applies_to"]:
                p5 = doc.add_paragraph(style="List Bullet")
                p5.paragraph_format.left_indent = Inches(0.4)
                run = p5.add_run(item)
                run.font.size = Pt(9.5)

            doc.add_paragraph()  # spacing

    # Undownloaded
    doc.add_heading("Citation-Only References (다운로드 불가)", level=1)
    p = doc.add_paragraph()
    add_inline_runs(p,
        "다음 문헌들은 paywall, 사이트 차단, 또는 서적이라 직접 다운로드되지 않았습니다. "
        "URL이 있으면 방문하여 접근 가능합니다.")

    for ref in UNDOWNLOADED_REFERENCES:
        p = doc.add_paragraph()
        run = p.add_run(f"• {ref['title']}")
        run.bold = True
        run.font.size = Pt(10)

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.2)
        run = p2.add_run(f"{ref['authors']} ({ref['year']}). {ref['venue']}.")
        run.italic = True
        run.font.size = Pt(9)

        if ref.get("url") and ref["url"] != "—":
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Inches(0.2)
            run = p3.add_run(f"URL: {ref['url']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)

        p4 = doc.add_paragraph()
        p4.paragraph_format.left_indent = Inches(0.2)
        run = p4.add_run(f"Applies: {ref['applies_to']}")
        run.font.size = Pt(9)

        if ref.get("note"):
            p5 = doc.add_paragraph()
            p5.paragraph_format.left_indent = Inches(0.2)
            run = p5.add_run(f"Note: {ref['note']}")
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x99, 0x6B, 0x00)
        doc.add_paragraph()

    # Footer / Info
    doc.add_heading("Folder Structure", level=1)
    folder_info = doc.add_paragraph()
    add_inline_runs(folder_info,
        "모든 PDF는 `docs/references/` 폴더에 저장되어 있습니다. "
        "파일명은 `[ID]_[Authors]_[Year]_[ShortTitle].pdf` 형식입니다.")

    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


if __name__ == "__main__":
    make_doc()
